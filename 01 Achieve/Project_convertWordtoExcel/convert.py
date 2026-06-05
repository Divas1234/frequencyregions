#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Word/HTML 表格提取与 Excel 转换工具
功能：自动识别输入 Word/HTML 文件，提取表格，保留单元格合并，下载并嵌入图片，生成排版精美的 Excel 报表。
"""

import os
import sys
import re
import argparse
import shutil
import tempfile
from io import BytesIO
import requests
from bs4 import BeautifulSoup
from PIL import Image as PILImage

# 导入 openpyxl 样式与处理模块
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.drawing.image import Image as OpenpyxlImage
from openpyxl.utils import get_column_letter

# 尝试导入 python-docx 用于支持真实 docx 解析
try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def check_file_type(file_path):
    """
    通过读取文件头部字节（Magic Bytes）和内容特征来探测文件真实类型
    返回：'html', 'docx', 'binary_doc', 'unknown'
    """
    if not os.path.exists(file_path):
        return 'unknown'
        
    try:
        with open(file_path, 'rb') as f:
            header = f.read(1024)  # 读取前1KB内容
    except Exception as e:
        print(f"[-] 无法读取文件: {e}")
        return 'unknown'

    # 1. 检查是否为 HTML 格式
    # 很多系统导出的 .doc 实际上是 HTML 格式，通常包含 <html> 或 <table 标签
    header_str = ""
    try:
        header_str = header.decode('utf-8', errors='ignore').lower().strip()
    except Exception:
        pass

    if '<html>' in header_str or '<html' in header_str or '<table' in header_str or '<!doctype html' in header_str:
        return 'html'

    # 2. 检查是否为 ZIP 压缩格式 (DOCX 文件的 Magic Number 是 PK\x03\x04)
    if header.startswith(b'PK\x03\x04'):
        return 'docx'

    # 3. 检查是否为传统二进制 DOC 格式 (Magic Number 是 \xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1)
    if header.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'):
        return 'binary_doc'

    # 如果以上都不是，做更深入的 HTML 标签匹配（以防编码格式不同）
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read(5000).lower()
            if '<table' in content or '<html' in content:
                return 'html'
    except Exception:
        pass

    return 'unknown'


def download_and_merge_images(img_urls, temp_dir, cell_coord):
    """
    下载网络图片，缩放并拼合成单张图片，避免在 Excel 单元格中发生重叠
    """
    downloaded_paths = []
    
    # 1. 下载所有图片
    for idx, url in enumerate(img_urls):
        # 兼容相对路径和没有协议头的链接
        if url.startswith('//'):
            url = 'https:' + url
        elif not url.startswith('http://') and not url.startswith('https://'):
            # 如果是本地路径，直接检查是否存在
            if os.path.exists(url):
                downloaded_paths.append(url)
                continue
            else:
                print(f"[!] 跳过无效图片路径: {url}")
                continue

        try:
            # 下载图片
            headers = {
                "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
            }
            response = requests.get(url, headers=headers, timeout=15)
            if response.status_code == 200:
                img_data = response.content
                # 写入临时文件
                ext = '.jpg'
                if 'image/png' in response.headers.get('Content-Type', ''):
                    ext = '.png'
                elif 'image/gif' in response.headers.get('Content-Type', ''):
                    ext = '.gif'
                
                temp_file_path = os.path.join(temp_dir, f"{cell_coord}_{idx}{ext}")
                with open(temp_file_path, 'wb') as f:
                    f.write(img_data)
                downloaded_paths.append(temp_file_path)
            else:
                print(f"[!] 下载图片失败 (HTTP {response.status_code}): {url}")
        except Exception as e:
            print(f"[!] 下载图片异常 ({url}): {e}")

    if not downloaded_paths:
        return None

    # 2. 将多张图片拼接成一张大图 (水平拼接，保证 Excel 的排版更合理)
    try:
        opened_images = []
        for path in downloaded_paths:
            try:
                img = PILImage.open(path)
                # 转换色彩模式，防止保存为 PNG/JPG 时报错
                if img.mode not in ('RGB', 'RGBA'):
                    img = img.convert('RGB')
                opened_images.append(img)
            except Exception as e:
                print(f"[!] 无法用 PIL 打开图片 {path}: {e}")

        if not opened_images:
            return None

        # 缩放所有图片到统一高度 (例如最大高度 100 像素)
        target_height = 100
        resized_images = []
        for img in opened_images:
            w, h = img.size
            scale = target_height / h
            new_w = max(1, int(w * scale))
            new_h = target_height
            # 限制单张图片最大宽度为 150 像素，防止列宽过大
            if new_w > 150:
                new_w = 150
                new_h = max(1, int(h * (150 / w)))
            
            resized_images.append(img.resize((new_w, new_h), PILImage.Resampling.LANCZOS))

        # 横向拼接图片，并在图片之间留有 5 像素的缝隙
        gap = 5
        total_width = sum(img.width for img in resized_images) + gap * (len(resized_images) - 1)
        max_height = max(img.height for img in resized_images)

        # 创建白色背景画布
        combined_img = PILImage.new('RGB', (total_width, max_height), color=(255, 255, 255))
        
        x_offset = 0
        for img in resized_images:
            # 垂直居中贴图
            y_offset = (max_height - img.height) // 2
            combined_img.paste(img, (x_offset, y_offset))
            x_offset += img.width + gap

        # 将拼接后的图片保存到临时文件中
        combined_path = os.path.join(temp_dir, f"{cell_coord}_combined.png")
        combined_img.save(combined_path, "PNG")
        return combined_path
    except Exception as e:
        print(f"[!] 拼接图片失败: {e}")
        return downloaded_paths[0] if downloaded_paths else None


def parse_html_tables(html_path):
    """
    解析 HTML 格式表格，支持 colspan/rowspan 以及图片提取
    """
    try:
        with open(html_path, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = f.read()
    except Exception as e:
        print(f"[-] 无法读取 HTML 文件: {e}")
        return []

    soup = BeautifulSoup(html_content, 'html.parser')
    tables = soup.find_all('table')
    if not tables:
        # 如果没找到 table 标签，尝试从 body 提取
        print("[!] HTML 中未找到 <table> 标签")
        return []

    print(f"[+] 找到 {len(tables)} 个 HTML 表格")
    parsed_tables = []

    for t_idx, table in enumerate(tables):
        rows = table.find_all('tr')
        if not rows:
            continue

        # 使用坐标字典记录所有单元格状态: (r_idx, c_idx) -> cell_info
        grid = {}
        max_cols = 0
        r_idx = 0

        for tr in rows:
            cells = tr.find_all(['td', 'th'])
            c_idx = 0
            for cell in cells:
                # 寻找当前行中第一个未被占用的列位置（可能被上一行的 rowspan 占用了）
                while (r_idx, c_idx) in grid:
                    c_idx += 1

                rowspan = int(cell.get('rowspan', 1))
                colspan = int(cell.get('colspan', 1))

                # 提取文本内容并去除两端空格
                text = cell.get_text(separator='\n', strip=True)

                # 提取图片 URL
                img_urls = []
                for img in cell.find_all('img'):
                    src = img.get('src')
                    if src:
                        img_urls.append(src)

                is_header = (cell.name == 'th' or 
                             tr.parent.name == 'thead' or 
                             cell.get('bgcolor') or 
                             'header' in str(cell.get('class', '')).lower())

                # 填充该单元格跨越的网格区域
                for dr in range(rowspan):
                    for dc in range(colspan):
                        curr_r = r_idx + dr
                        curr_c = c_idx + dc
                        is_origin = (dr == 0 and dc == 0)
                        grid[(curr_r, curr_c)] = {
                            'text': text if is_origin else "",
                            'img_urls': img_urls if is_origin else [],
                            'is_origin': is_origin,
                            'origin_pos': (r_idx, c_idx),
                            'span': (rowspan, colspan),
                            'is_header': is_header
                        }

                c_idx += colspan
                max_cols = max(max_cols, c_idx)
            r_idx += 1

        max_rows = r_idx
        parsed_tables.append({
            'index': t_idx,
            'grid': grid,
            'max_rows': max_rows,
            'max_cols': max_cols
        })
        print(f"    - 表格 #{t_idx+1}: 规格为 {max_rows} 行 x {max_cols} 列")

    return parsed_tables


def parse_docx_tables(docx_path):
    """
    解析标准的二进制 DOCX 表格，支持跨行/跨列合并
    """
    if not HAS_DOCX:
        print("[-] 缺少 python-docx 依赖，无法解析标准的 .docx 文件！请先安装依赖。")
        return []

    try:
        doc = docx.Document(docx_path)
    except Exception as e:
        print(f"[-] 无法读取 DOCX 文件: {e}")
        return []

    print(f"[+] 找到 {len(doc.tables)} 个 DOCX 表格")
    parsed_tables = []

    for t_idx, table in enumerate(doc.tables):
        max_rows = len(table.rows)
        if max_rows == 0:
            continue
        max_cols = len(table.columns)

        grid = {}
        seen_cells = set()

        for r_idx in range(max_rows):
            for c_idx in range(max_cols):
                try:
                    cell = table.cell(r_idx, c_idx)
                except Exception:
                    continue

                # docx 中，合并的单元格共享同一个 XML 元素底层 (_tc)
                cell_id = cell._tc
                if cell_id in seen_cells:
                    # 已经被处理过的合并区域单元格，只保留占位
                    if (r_idx, c_idx) not in grid:
                        grid[(r_idx, c_idx)] = {
                            'text': "",
                            'img_urls': [],
                            'is_origin': False,
                            'origin_pos': None,
                            'span': (1, 1),
                            'is_header': False
                        }
                    continue

                seen_cells.add(cell_id)

                # 计算跨列 colspan
                colspan = 1
                for tc in range(c_idx + 1, max_cols):
                    try:
                        if table.cell(r_idx, tc)._tc == cell_id:
                            colspan += 1
                        else:
                            break
                    except Exception:
                        break

                # 计算跨行 rowspan
                rowspan = 1
                for tr in range(r_idx + 1, max_rows):
                    try:
                        if table.cell(tr, c_idx)._tc == cell_id:
                            rowspan += 1
                        else:
                            break
                    except Exception:
                        break

                text = cell.text.strip()
                is_header = r_idx == 0 # 默认首行为表头

                # 填充网格
                for dr in range(rowspan):
                    for dc in range(colspan):
                        curr_r = r_idx + dr
                        curr_c = c_idx + dc
                        is_origin = (dr == 0 and dc == 0)
                        grid[(curr_r, curr_c)] = {
                            'text': text if is_origin else "",
                            'img_urls': [], # docx 的内联图片提取由于格式复杂暂时跳过，若需要可以通过段落 XML 获取
                            'is_origin': is_origin,
                            'origin_pos': (r_idx, c_idx),
                            'span': (rowspan, colspan),
                            'is_header': is_header
                        }

        parsed_tables.append({
            'index': t_idx,
            'grid': grid,
            'max_rows': max_rows,
            'max_cols': max_cols
        })
        print(f"    - 表格 #{t_idx+1}: 规格为 {max_rows} 行 x {max_cols} 列")

    return parsed_tables


def export_tables_to_excel(parsed_tables, output_path, fill_merged=True, download_images=True):
    """
    将解析出的表格写入 Excel，包含排版美化、单元格合并与图片嵌入
    """
    # 创建临时文件夹来存储图片
    temp_dir = tempfile.mkdtemp(prefix="docx_to_excel_")
    
    try:
        wb = openpyxl.Workbook()
        # 默认会创建一个名为 "Sheet" 的表，我们直接重命名或删除
        default_sheet = wb.active
        
        # 定义高雅专业的样式
        # 1. 字体
        font_family = "微软雅黑"
        header_font = Font(name=font_family, size=11, bold=True, color="FFFFFF")
        data_font = Font(name=font_family, size=10, color="000000")
        note_font = Font(name=font_family, size=9, italic=True, color="555555")
        
        # 2. 填充颜色 (深石墨青色表头，典雅大方)
        header_fill = PatternFill(start_color="2F4F4F", end_color="2F4F4F", fill_type="solid")
        # 偶数行斑马线背景 (极淡的浅灰色)
        zebra_fill = PatternFill(start_color="F9FAFB", end_color="F9FAFB", fill_type="solid")
        
        # 3. 边框 (淡雅细灰线)
        thin_border_side = Side(style='thin', color='D3D3D3')
        thin_border = Border(left=thin_border_side, right=thin_border_side, top=thin_border_side, bottom=thin_border_side)
        
        # 4. 对齐方式
        center_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        left_alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)

        for table_idx, table_data in enumerate(parsed_tables):
            grid = table_data['grid']
            max_rows = table_data['max_rows']
            max_cols = table_data['max_cols']
            
            # 创建新 Worksheet
            sheet_title = f"数据表_{table_idx + 1}"
            if table_idx == 0:
                ws = default_sheet
                ws.title = sheet_title
            else:
                ws = wb.create_sheet(title=sheet_title)
            
            # 显示默认网格线
            ws.views.sheetView[0].showGridLines = True
            
            # 记录每列的最大字符宽度和图片宽度，以便最后自适应列宽
            col_widths = {c: 10 for c in range(1, max_cols + 1)}
            row_heights = {r: 24 for r in range(1, max_rows + 1)} # 默认行高 24pt
            
            # 1. 写入所有单元格数据与应用基本样式
            for r in range(max_rows):
                row_num = r + 1
                for c in range(max_cols):
                    col_num = c + 1
                    cell_key = (r, c)
                    
                    if cell_key not in grid:
                        continue
                        
                    cell_info = grid[cell_key]
                    excel_cell = ws.cell(row=row_num, column=col_num)
                    
                    is_origin = cell_info['is_origin']
                    origin_pos = cell_info['origin_pos']
                    
                    # 确定要写入的内容
                    val = ""
                    if is_origin:
                        val = cell_info['text']
                    elif fill_merged:
                        # 启用填充时，从合并源头单元格拷贝数据
                        if origin_pos and origin_pos in grid:
                            val = grid[origin_pos]['text']
                            
                    # 处理值类型，尝试转为数字，以防统计公式失效
                    if val.isdigit():
                        excel_cell.value = int(val)
                    else:
                        # 如果是小数
                        try:
                            excel_cell.value = float(val) if '.' in val else val
                        except ValueError:
                            excel_cell.value = val

                    # 基础样式应用
                    excel_cell.font = data_font
                    excel_cell.border = thin_border
                    
                    # 根据是否为表头应用样式
                    if cell_info['is_header']:
                        excel_cell.fill = header_fill
                        excel_cell.font = header_font
                        excel_cell.alignment = center_alignment
                    else:
                        # 对齐逻辑：短文本和特殊内容居中，长文本靠左
                        val_str = str(excel_cell.value or "")
                        if len(val_str) <= 6 or val_str.isdigit() or re.match(r'^\d{4}-\d{2}-\d{2}', val_str):
                            excel_cell.alignment = center_alignment
                        else:
                            excel_cell.alignment = left_alignment
                        
                        # 斑马线条纹
                        if r % 2 == 1:
                            excel_cell.fill = zebra_fill

                    # 2. 统计文本宽度，计算列宽
                    if val:
                        lines = str(val).split('\n')
                        max_line_len = max(len(l.encode('gbk', errors='ignore')) for l in lines)
                        # 中文字符宽度折算：gbk 编码长度即是显示宽度
                        col_widths[col_num] = max(col_widths[col_num], max_line_len + 3)
                    
                    # 3. 处理图片嵌入
                    img_urls = cell_info['img_urls']
                    if download_images and img_urls and is_origin:
                        cell_coord = f"{get_column_letter(col_num)}{row_num}"
                        print(f"[*] 正在处理单元格 {cell_coord} 的图片下载...")
                        
                        combined_path = download_and_merge_images(img_urls, temp_dir, cell_coord)
                        if combined_path:
                            try:
                                # 使用 openpyxl 包装图片并插入
                                img = OpenpyxlImage(combined_path)
                                # 写入图片
                                ws.add_image(img, cell_coord)
                                
                                # 将该行行高设为图片高度 + 间距（Excel行高 1pt ≈ 1.33px）
                                img_height_pt = int(img.height / 1.33) + 15
                                row_heights[row_num] = max(row_heights[row_num], img_height_pt)
                                
                                # 更新图片所在列的宽度（Excel 1单位宽度 ≈ 7.2px）
                                img_width_units = int(img.width / 7.2) + 4
                                col_widths[col_num] = max(col_widths[col_num], img_width_units)
                                
                                # 清空文本值，避免文字挡住图片，或仅保留隐患数指示
                                if not excel_cell.value:
                                    excel_cell.value = " " # 用空格占位
                            except Exception as img_err:
                                print(f"[-] 图片插入 Excel 失败: {img_err}")
                                # 降级为保留链接
                                excel_cell.value = "\n".join(img_urls)
                        else:
                            # 降级：如果图片下载失败，保留链接文本
                            excel_cell.value = "\n".join(img_urls)
                            excel_cell.font = note_font

            # 4. 执行单元格合并 (仅在没有启用 fill_merged 时)
            if not fill_merged:
                merged_ranges = set()
                for r in range(max_rows):
                    for c in range(max_cols):
                        cell_key = (r, c)
                        if cell_key not in grid:
                            continue
                        
                        cell_info = grid[cell_key]
                        if cell_info['is_origin']:
                            rowspan, colspan = cell_info['span']
                            if rowspan > 1 or colspan > 1:
                                start_r = r + 1
                                start_c = c + 1
                                end_r = r + rowspan
                                end_c = c + colspan
                                
                                merge_range = (start_r, start_c, end_r, end_c)
                                if merge_range not in merged_ranges:
                                    ws.merge_cells(start_row=start_r, start_column=start_c, 
                                                   end_row=end_r, end_column=end_c)
                                    merged_ranges.add(merge_range)
                                    
                                    # 修正合并后非左上角单元格的边框展示 bug
                                    # openpyxl 在合并后，除了左上角，其他格子的边框默认会消失。需要手动补齐 border。
                                    for mr in range(start_r, end_r + 1):
                                        for mc in range(start_c, end_c + 1):
                                            ws.cell(row=mr, column=mc).border = thin_border

            # 5. 设置每列自适应宽度与行高
            for col_num, width in col_widths.items():
                # 限制最大列宽为 55 字符，防止某些长描述列过宽
                final_width = min(width, 55)
                ws.column_dimensions[get_column_letter(col_num)].width = final_width
                
            for row_num, height in row_heights.items():
                ws.row_dimensions[row_num].height = height

        # 保存 Excel
        wb.save(output_path)
        print(f"[+] 转换成功！Excel 文件已保存至: {output_path}")
        
    finally:
        # 清理临时下载的图片文件
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)


def main():
    parser = argparse.ArgumentParser(description="提取 Word 表格并转化为精美的 Excel 格式 (.xlsx)")
    parser.add_argument("input", help="输入的 Word/HTML 文件路径（或文件夹以执行批量转换）")
    parser.add_argument("-o", "--output", help="输出的 Excel 文件路径 (默认同目录下同名 .xlsx)")
    parser.add_argument("--keep-merged", action="store_true", help="保留单元格的合并样式（默认拆分并填充，方便统计）")
    parser.add_argument("--fill-merged", action="store_true", help="填充合并单元格的内容（已弃用，现在默认开启）")
    parser.add_argument("--no-images", action="store_true", help="禁用提取和下载图片到 Excel 中")
    
    args = parser.parse_args()
    
    # 默认启用拆分与填充。如果指定了 --keep-merged，则保留单元格的合并样式。
    fill_merged = not args.keep_merged
    
    input_path = os.path.abspath(args.input)
    
    # 批量处理目录模式
    if os.path.isdir(input_path):
        print(f"[*] 检测到输入为文件夹，开始批量扫描: {input_path}")
        files = [os.path.join(input_path, f) for f in os.listdir(input_path) 
                 if f.endswith(('.doc', '.docx', '.html', '.htm')) and not f.startswith('~$')]
        
        if not files:
            print("[-] 未在指定文件夹内找到支持的文档格式（.doc, .docx, .html）")
            sys.exit(1)
            
        print(f"[+] 共扫描到 {len(files)} 个待转换的文件。")
        for idx, file_file in enumerate(files):
            print(f"\n[{idx+1}/{len(files)}] 正在处理: {os.path.basename(file_file)}")
            file_type = check_file_type(file_file)
            print(f"[*] 探测到文件真实类型: {file_type.upper()}")
            
            # 生成默认输出路径
            base, _ = os.path.splitext(file_file)
            out_file = base + ".xlsx"
            
            # 执行解析
            tables = []
            if file_type == 'html':
                tables = parse_html_tables(file_file)
            elif file_type == 'docx':
                tables = parse_docx_tables(file_file)
            else:
                print(f"[-] 暂不支持该类型的提取，跳过: {os.path.basename(file_file)}")
                continue
                
            if not tables:
                print("[-] 文件内未提取到有效表格，跳过转换。")
                continue
                
            export_tables_to_excel(tables, out_file, fill_merged=fill_merged, download_images=not args.no_images)
            
        print("\n[+] 批量转换全部完成！")
        sys.exit(0)

    # 单个文件处理模式
    if not os.path.exists(input_path):
        print(f"[-] 输入的文件路径不存在: {args.input}")
        sys.exit(1)

    print(f"[*] 正在分析文件: {args.input}")
    file_type = check_file_type(input_path)
    print(f"[*] 文件真实格式为: {file_type.upper()}")

    # 确定输出路径
    if args.output:
        output_path = os.path.abspath(args.output)
    else:
        base, _ = os.path.splitext(input_path)
        output_path = base + ".xlsx"

    # 根据类型选择解析引擎
    tables = []
    if file_type == 'html':
        tables = parse_html_tables(input_path)
    elif file_type == 'docx':
        tables = parse_docx_tables(input_path)
    elif file_type == 'binary_doc':
        print("\n[-] 探测到为传统微软二进制安全 DOC 格式（非 HTML 包装）。")
        print("[-] 在 macOS 环境下，Python 无法直接解析此二进制格式。")
        print("[!] 建议手动在 Microsoft Word 中将该文件【另存为】.docx 或网页 (.html) 格式，然后再重新运行此脚本进行转换。")
        sys.exit(1)
    else:
        print("[-] 无法识别文件内容格式，解析失败。")
        sys.exit(1)

    if not tables:
        print("[-] 文档中未找到任何表格。")
        sys.exit(1)

    # 执行导出
    export_tables_to_excel(tables, output_path, fill_merged=fill_merged, download_images=not args.no_images)


if __name__ == "__main__":
    main()
